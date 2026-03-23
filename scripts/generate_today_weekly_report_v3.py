from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd
import pyarrow.parquet as pq
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = Path(
    r"C:\Users\19672\xwechat_files\wxid_l9yi1ctvtxw122_fa84\msg\file\2026-03\9th_Weekly_Report_v2.docx"
)
OUTPUT_DIR = ROOT / "output" / "doc"
DOCX_PATH = OUTPUT_DIR / "11th_Weekly_Report_2026-03-23_Updated.docx"
MARKDOWN_PATH = OUTPUT_DIR / "11th_Weekly_Report_2026-03-23_Updated.md"

SUITE_DIR = ROOT / "outputs" / "report_version_suite" / "report_version_suite_current_data_20260323"
FLIP_DIR = ROOT / "outputs" / "analysis" / "report_version_suite_current_data_20260323_v1_flip"
RUNS_DIR = ROOT / "outputs" / "runs"
V1_RUN_DIR = RUNS_DIR / "report_version_suite_current_data_20260323_v1"


@dataclass
class SplitStats:
    split: str
    rows: int
    permnos: int
    positive_rate: float


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class Paragraph:
    text: str


@dataclass
class BulletList:
    items: list[str]


@dataclass
class TableBlock:
    headers: list[str]
    rows: list[list[str]]


Block = Heading | Paragraph | BulletList | TableBlock


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def pct(value: float, digits: int = 2) -> str:
    return f"{value * 100:.{digits}f}%"


def num(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def format_bp(value: float) -> str:
    return f"{value * 10000:.1f} bp"


def load_split_stats() -> tuple[list[SplitStats], int, int, float]:
    stats: list[SplitStats] = []
    all_permnos: set[int] = set()
    total_rows = 0
    total_positive = 0.0
    for split in ["train", "val", "test"]:
        table = pq.read_table(ROOT / "data" / "processed" / f"{split}.parquet", columns=["PERMNO", "y_div_10d"])
        df = table.to_pandas()
        rows = len(df)
        permnos = int(df["PERMNO"].nunique())
        positive_rate = float(df["y_div_10d"].mean())
        stats.append(SplitStats(split=split, rows=rows, permnos=permnos, positive_rate=positive_rate))
        all_permnos.update(df["PERMNO"].astype(int).tolist())
        total_rows += rows
        total_positive += rows * positive_rate
    total_permnos = len(all_permnos)
    total_positive_rate = total_positive / total_rows if total_rows else 0.0
    return stats, total_rows, total_permnos, total_positive_rate


def yearly_summary(yearly_compare: pd.DataFrame, other: str) -> tuple[int, int, int, float, int, float]:
    gap_col = f"excess_annualized_return_v1_minus_{other}"
    valid = yearly_compare[["year", gap_col]].dropna()
    wins = int((valid[gap_col] > 0).sum())
    total = int(len(valid))
    best = valid.loc[valid[gap_col].idxmax()]
    worst = valid.loc[valid[gap_col].idxmin()]
    return wins, total, int(best["year"]), float(best[gap_col]), int(worst["year"]), float(worst[gap_col])


def top_industries(industry_compare: pd.DataFrame, other: str, top_n: int = 3) -> tuple[list[str], list[str]]:
    gap_col = f"mean_trade_return_v1_minus_{other}"
    ranked = industry_compare[["industry", gap_col]].dropna().sort_values(gap_col, ascending=False)
    best = [f"{int(row.industry)} ({pct(float(getattr(row, gap_col)))})" for row in ranked.head(top_n).itertuples(index=False)]
    worst = [f"{int(row.industry)} ({pct(float(getattr(row, gap_col)))})" for row in ranked.tail(top_n).sort_values(gap_col).itertuples(index=False)]
    return best, worst


def safe_style(doc: Document, preferred: str, fallback: str) -> str:
    return preferred if preferred in doc.styles else fallback


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.88)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    bullet_style = safe_style(doc, "List Bullet", "List Paragraph")
    for item in items:
        p = doc.add_paragraph(style=bullet_style)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if bullet_style == "List Paragraph":
            p.add_run(f"- {item}")
        else:
            p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid" if "Table Grid" in doc.styles else "Normal Table"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        for paragraph in hdr_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def render_docx(blocks: list[Block]) -> None:
    doc = Document(TEMPLATE_PATH)
    doc._body.clear_content()

    title = doc.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("11th Weekly Report Submission")
    run.bold = True

    for line in [
        "Project: Corporate Action Prediction Model Based on Multi-Source Data",
        "Group Members: Zuohong Pan (潘祚弘), Wenjing Liu (刘文菁), Zheyu Hu (胡哲越)",
        "Week: 11 (2026-03-23)",
    ]:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(line)

    for block in blocks:
        if isinstance(block, Heading):
            style = f"Heading {block.level}"
            doc.add_paragraph(block.text, style=style)
        elif isinstance(block, Paragraph):
            add_body_paragraph(doc, block.text)
        elif isinstance(block, BulletList):
            add_bullets(doc, block.items)
        elif isinstance(block, TableBlock):
            add_table(doc, block.headers, block.rows)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def render_markdown(blocks: list[Block]) -> str:
    lines: list[str] = [
        "# 11th Weekly Report Submission",
        "",
        "Project: Corporate Action Prediction Model Based on Multi-Source Data",
        "Group Members: Zuohong Pan (潘祚弘), Wenjing Liu (刘文菁), Zheyu Hu (胡哲越)",
        "Week: 11 (2026-03-23)",
        "",
    ]
    for block in blocks:
        if isinstance(block, Heading):
            lines.append(f"{'#' * (block.level + 1)} {block.text}")
            lines.append("")
        elif isinstance(block, Paragraph):
            lines.append(block.text)
            lines.append("")
        elif isinstance(block, BulletList):
            for item in block.items:
                lines.append(f"- {item}")
            lines.append("")
        elif isinstance(block, TableBlock):
            lines.append("| " + " | ".join(block.headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(block.headers)) + " |")
            for row in block.rows:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
    return "\n".join(lines)


def build_blocks() -> list[Block]:
    split_stats, total_rows, total_permnos, total_positive_rate = load_split_stats()
    suite_df = pd.read_csv(SUITE_DIR / "suite_results.csv")
    suite_df = suite_df.sort_values(["test_excess_annualized_return", "test_portfolio_sharpe"], ascending=[False, False]).reset_index(drop=True)
    version_map = {row["version"]: row for _, row in suite_df.iterrows()}

    yearly_compare = pd.read_csv(FLIP_DIR / "yearly_backtest_compare.csv")
    history_compare = pd.read_csv(FLIP_DIR / "history_trade_compare.csv")
    liquidity_compare = pd.read_csv(FLIP_DIR / "liquidity_trade_compare.csv")
    industry_compare = pd.read_csv(FLIP_DIR / "industry_trade_compare.csv")

    v1_metrics = read_json(V1_RUN_DIR / "standardized" / "metrics.json")
    v7_metrics = read_json(RUNS_DIR / "report_version_suite_current_data_20260323_v7" / "standardized" / "metrics.json")
    v8_metrics = read_json(RUNS_DIR / "report_version_suite_current_data_20260323_v8" / "standardized" / "metrics.json")
    summary = read_json(V1_RUN_DIR / "backtest" / "summary.json")
    ranking = read_json(V1_RUN_DIR / "backtest" / "reference" / "ranking_comparison.json")
    three_way = read_json(V1_RUN_DIR / "backtest" / "reference" / "three_way_comparison.json")

    v1_vs_v7 = yearly_summary(yearly_compare, "v7")
    v1_vs_v8 = yearly_summary(yearly_compare, "v8")
    history_row = history_compare.iloc[0]
    liquidity_rows = liquidity_compare[liquidity_compare["liquidity_bucket"] != "unknown"].copy()
    v7_best, v7_worst = top_industries(industry_compare, "v7")
    v8_best, v8_worst = top_industries(industry_compare, "v8")

    ranking_metrics = ranking["metrics"]
    ranking_deltas = ranking["deltas"]
    strategy_vs_event = three_way["strategy_vs_event_oracle"]

    blocks: list[Block] = []

    blocks.extend(
        [
            Heading(1, "1. 本周工作概述"),
            Paragraph(
                "本周工作围绕四条主线展开：其一，在当前正确数据版本下重跑 v1~v8 与 random 全版本实验，确认旧中期报告中的“v7 最优”结论不再成立，当前最优基线已经回到 v1；其二，针对 v1 为何反超 v7/v8 做分层诊断，按年份、分红历史、流动性与行业拆解收益与命中率差异；其三，重构回测参考系，补齐 non_prob、random_prob、prob 与 y_div_10d 四路排序比较，同时澄清 oracle_event 与 oracle_return 的口径边界；其四，固化实验版本与标准化评估管线，把版本注册、统一配置、标准输出和测试用例整理进正式工程入口。"
            ),
            Paragraph(
                f"在当前数据版本下，v1 测试期组合年化收益为 {pct(summary['portfolio']['annualized_return'])}，相对基准超额年化收益为 {pct(summary['excess_vs_benchmark']['annualized_return'])}，Sharpe 为 {num(summary['portfolio']['sharpe'])}；同时 auto policy 在 held-out 验证口径下得到 {summary['research_decision']['support_votes']}/{summary['research_decision']['total_checks']} 支持票并启用分红规则。"
            ),
            BulletList(
                [
                    "完成当前数据版本的 v1~v8 与 random 全版本复现实验，并重写配套实验报告。",
                    "完成 v1 相对 v7/v8 的分层翻转分析，明确优势集中在 has_history 与中等流动性样本。",
                    "完成四路排序对比：non_prob、random_prob、prob、y_div_10d。",
                    "完成版本注册表、正式配置目录、标准化评估入口和回归测试扩展。",
                ]
            ),
            Heading(1, "2. 当前数据版本下的版本复现实验结果"),
            Paragraph(
                "当前使用的数据面板覆盖 train、val、test 三个切分，目标标签仍为未来 10 个交易日内是否发生分红公告。面板总体规模与旧报告相比已经发生变化，因此旧报告中的最优版本排序不能直接沿用。"
            ),
            TableBlock(
                headers=["Split", "Rows", "Unique PERMNO", "Positive Rate"],
                rows=[
                    [stat.split, f"{stat.rows:,}", str(stat.permnos), pct(stat.positive_rate)]
                    for stat in split_stats
                ]
                + [["total", f"{total_rows:,}", str(total_permnos), pct(total_positive_rate)]],
            ),
            Paragraph(
                "在同一数据版本、同一回测口径下比较各版本后，测试期超额收益和 Sharpe 的排序已经明显改变。v1 虽然不是验证集 AUPRC 最高的模型，但在真实回测上反而取得了当前最强的净效果。"
            ),
            TableBlock(
                headers=["Version", "Val AUPRC", "Test Ann", "Test Excess Ann", "Sharpe", "Policy Votes", "Use Rules"],
                rows=[
                    [
                        str(row["version"]),
                        pct(float(row["val_aucpr"])),
                        pct(float(row["test_portfolio_annualized_return"])),
                        pct(float(row["test_excess_annualized_return"])),
                        num(float(row["test_portfolio_sharpe"])),
                        f"{int(row['policy_support_votes'])}/{int(row['policy_total_checks'])}",
                        "是" if bool(row["policy_use_dividend_rules"]) else "否",
                    ]
                    for _, row in suite_df.iterrows()
                ],
            ),
            BulletList(
                [
                    f"v1 当前是最优基线，测试期超额年化收益为 {pct(float(version_map['v1']['test_excess_annualized_return']))}，Sharpe 为 {num(float(version_map['v1']['test_portfolio_sharpe']))}。",
                    f"v7 的验证集 AUPRC 更高，达到 {pct(float(version_map['v7']['val_aucpr']))}，但测试期超额年化收益仅为 {pct(float(version_map['v7']['test_excess_annualized_return']))}。",
                    f"v8 在当前数据上不再像旧报告那样“完全失效”，但测试期超额年化收益仍为 {pct(float(version_map['v8']['test_excess_annualized_return']))}，低于 v1。",
                    f"random 端到端基线的验证集 AUPRC 只有 {pct(float(version_map['random']['val_aucpr']))}，测试期超额年化收益为 {pct(float(version_map['random']['test_excess_annualized_return']))}，显著弱于任一有效版本。",
                ]
            ),
            Heading(1, "3. v1 反超 v7/v8 的分层分析"),
            Heading(2, "3.1 年度维度"),
            TableBlock(
                headers=["Comparison", "Winning Years", "Total Years", "Best Year", "Best Gap", "Worst Year", "Worst Gap"],
                rows=[
                    ["v1 vs v7", str(v1_vs_v7[0]), str(v1_vs_v7[1]), str(v1_vs_v7[2]), pct(v1_vs_v7[3]), str(v1_vs_v7[4]), pct(v1_vs_v7[5])],
                    ["v1 vs v8", str(v1_vs_v8[0]), str(v1_vs_v8[1]), str(v1_vs_v8[2]), pct(v1_vs_v8[3]), str(v1_vs_v8[4]), pct(v1_vs_v8[5])],
                ],
            ),
            Paragraph(
                "按年份拆开看，v1 相比 v7 在 10 年中赢了 9 年，相比 v8 赢了 7 年。最大优势都出现在 2019 年，而 2021 年是 v1 相对落后的主要年份。这说明 v1 的领先不是由单一年度偶然驱动，而是更接近跨年份的稳定优势。"
            ),
            Heading(2, "3.2 分红历史与流动性维度"),
            TableBlock(
                headers=["History Bucket", "v1 Mean Trade Ret", "v1 Hit Rate", "v1-v7 Mean Ret Gap", "v1-v7 Hit Rate Gap", "v1-v8 Mean Ret Gap", "v1-v8 Hit Rate Gap"],
                rows=[
                    [
                        str(history_row["history_bucket"]),
                        pct(float(history_row["mean_trade_return_v1"])),
                        pct(float(history_row["hit_rate_v1"])),
                        pct(float(history_row["mean_trade_return_v1_minus_v7"])),
                        pct(float(history_row["hit_rate_v1_minus_v7"])),
                        pct(float(history_row["mean_trade_return_v1_minus_v8"])),
                        pct(float(history_row["hit_rate_v1_minus_v8"])),
                    ]
                ],
            ),
            Paragraph(
                "历史维度显示，当前实际成交几乎全部来自 has_history 子样本。就在这一组里，v1 的平均单笔收益比 v7 和 v8 都高出约 10bp；其中 v1 相对 v7 的事件命中率反而低 1.45 个百分点，说明 v1 的优势主要不是“更会猜中事件”，而是“即便命中率不占优，挑出来的交易更赚钱”。"
            ),
            TableBlock(
                headers=["Liquidity Bucket", "v1-v7 Mean Ret Gap", "v1-v7 Hit Rate Gap", "v1-v8 Mean Ret Gap", "v1-v8 Hit Rate Gap"],
                rows=[
                    [
                        str(row["liquidity_bucket"]),
                        pct(float(row["mean_trade_return_v1_minus_v7"])),
                        pct(float(row["hit_rate_v1_minus_v7"])),
                        pct(float(row["mean_trade_return_v1_minus_v8"])),
                        pct(float(row["hit_rate_v1_minus_v8"])),
                    ]
                    for _, row in liquidity_rows.iterrows()
                ],
            ),
            Paragraph(
                "流动性维度上，v1 的优势主要集中在 liq_q2 和 liq_q3 两档中等流动性样本，其中相对 v7 的平均单笔收益差分别约为 18bp 和 29bp。最高流动性的 liq_q4 并不是 v1 优势最强的来源，这说明 v1 更像是在中段流动性样本上学到了更有效的净收益排序。"
            ),
            Heading(2, "3.3 行业维度"),
            TableBlock(
                headers=["Comparison", "Top Positive Industries", "Top Negative Industries"],
                rows=[
                    ["v1 vs v7", ", ".join(v7_best), ", ".join(v7_worst)],
                    ["v1 vs v8", ", ".join(v8_best), ", ".join(v8_worst)],
                ],
            ),
            Paragraph(
                "行业维度的差异是分散的，而不是某一个单一行业决定全部结果。它更像是多个行业中同时存在中等幅度的净收益改善，因此“v1 优于 v7/v8”更适合表述为广泛分布的收益兑现差异，而不是简单归因给某一行业变量。"
            ),
            Heading(1, "4. 排序价值拆解：non_prob / random_prob / prob / y_div_10d"),
            Paragraph(
                "为了把规则价值、候选池价值与排序价值拆开，本周在共享同一套交易规则、流动性过滤、行业约束、持仓期与成本模型的前提下，补充了四条可直接对比的参考线：non_prob、random_prob、prob 和 y_div_10d。"
            ),
            TableBlock(
                headers=["Variant", "Annualized Return", "Sharpe", "Total Return", "Max Drawdown"],
                rows=[
                    ["non_prob", pct(float(ranking_metrics["non_prob"]["annualized_return"])), num(float(ranking_metrics["non_prob"]["sharpe"])), pct(float(ranking_metrics["non_prob"]["total_return"])), pct(float(ranking_metrics["non_prob"]["max_drawdown"]))],
                    ["random_prob", pct(float(ranking_metrics["random_prob"]["annualized_return"])), num(float(ranking_metrics["random_prob"]["sharpe"])), pct(float(ranking_metrics["random_prob"]["total_return"])), pct(float(ranking_metrics["random_prob"]["max_drawdown"]))],
                    ["prob", pct(float(ranking_metrics["prob"]["annualized_return"])), num(float(ranking_metrics["prob"]["sharpe"])), pct(float(ranking_metrics["prob"]["total_return"])), pct(float(ranking_metrics["prob"]["max_drawdown"]))],
                    ["y_div_10d", pct(float(ranking_metrics["y_div_10d"]["annualized_return"])), num(float(ranking_metrics["y_div_10d"]["sharpe"])), pct(float(ranking_metrics["y_div_10d"]["total_return"])), pct(float(ranking_metrics["y_div_10d"]["max_drawdown"]))],
                ],
            ),
            TableBlock(
                headers=["Delta", "Annualized Return Gap", "Sharpe Gap", "Total Return Gap"],
                rows=[
                    ["random_prob - non_prob", pct(float(ranking_deltas["random_prob_minus_non_prob"]["annualized_return_diff"])), num(float(ranking_deltas["random_prob_minus_non_prob"]["sharpe_diff"])), pct(float(ranking_deltas["random_prob_minus_non_prob"]["total_return_diff"]))],
                    ["prob - random_prob", pct(float(ranking_deltas["prob_minus_random_prob"]["annualized_return_diff"])), num(float(ranking_deltas["prob_minus_random_prob"]["sharpe_diff"])), pct(float(ranking_deltas["prob_minus_random_prob"]["total_return_diff"]))],
                    ["y_div_10d - prob", pct(float(ranking_deltas["y_div_10d_minus_prob"]["annualized_return_diff"])), num(float(ranking_deltas["y_div_10d_minus_prob"]["sharpe_diff"])), pct(float(ranking_deltas["y_div_10d_minus_prob"]["total_return_diff"]))],
                ],
            ),
            BulletList(
                [
                    f"non_prob 年化收益为 {pct(float(ranking_metrics['non_prob']['annualized_return']))}，说明仅靠 stable/regular、policy、流动性和行业规则本身，不足以形成稳定正收益。",
                    f"random_prob 年化收益升至 {pct(float(ranking_metrics['random_prob']['annualized_return']))}，相对 non_prob 提升 {pct(float(ranking_deltas['random_prob_minus_non_prob']['annualized_return_diff']))}，说明 prob 过滤后的候选池本身已经有显著信息价值。",
                    f"prob 相对 random_prob 再提升 {pct(float(ranking_deltas['prob_minus_random_prob']['annualized_return_diff']))}，证明模型排序本身仍贡献了额外 alpha。",
                    f"y_div_10d 年化收益低于 prob {pct(abs(float(ranking_deltas['y_div_10d_minus_prob']['annualized_return_diff'])))}，说明“知道事件是否会发生”并不等价于“知道哪笔交易更赚钱”。",
                ]
            ),
            Heading(1, "5. Oracle 基准口径重构与解释"),
            Paragraph(
                "本周同时梳理了 random baseline、oracle_event_ceiling 与 oracle_return_ceiling 三个参考基准的语义边界。关键点是：oracle_return_ceiling 假设知道未来持有期收益，属于诊断型的理论上界；oracle_event_ceiling 只共享策略过滤并按真实 y_div_10d 排序，是事件信息基准，而不是收益上界。"
            ),
            TableBlock(
                headers=["Reference", "Annualized Return", "Sharpe", "Total Return", "Definition"],
                rows=[
                    ["random_baseline", pct(float(three_way["random_baseline"]["annualized_return"])), num(float(three_way["random_baseline"]["sharpe"])), pct(float(three_way["random_baseline"]["total_return"])), "基础可交易过滤后的随机排序"],
                    ["strategy", pct(float(three_way["strategy"]["annualized_return"])), num(float(three_way["strategy"]["sharpe"])), pct(float(three_way["strategy"]["total_return"])), "实际策略结果"],
                    ["oracle_event_ceiling", pct(float(three_way["oracle_event_ceiling"]["annualized_return"])), num(float(three_way["oracle_event_ceiling"]["sharpe"])), pct(float(three_way["oracle_event_ceiling"]["total_return"])), "共享策略过滤后按真实 y_div_10d 排序"],
                    ["oracle_return_ceiling", pct(float(three_way["oracle_return_ceiling"]["annualized_return"])), num(float(three_way["oracle_return_ceiling"]["sharpe"])), pct(float(three_way["oracle_return_ceiling"]["total_return"])), "按未来持有期收益排序的诊断型上界"],
                ],
            ),
            Paragraph(
                f"当前 v1 的 strategy 年化收益为 {pct(float(three_way['strategy']['annualized_return']))}，比 oracle_event_ceiling 高 {pct(float(strategy_vs_event['annualized_return_diff']))}。这并不矛盾，因为 oracle_event_ceiling 只知道事件 hit，不知道同一批 hit 里哪只股票更赚钱；它是事件时点信息基准，而不是严格收益上界。与之相对，oracle_return_ceiling 的年化收益达到 {pct(float(three_way['oracle_return_ceiling']['annualized_return']))}，正是因为它额外假设了完美未来价格信息，因此不能被直接拿来评价策略是否“合理”。"
            ),
            Paragraph(
                f"从 alpha capture 角度看，当前 v1 的 primary capture ratio 仅为 {num(float(three_way['alpha_capture']['primary_alpha_capture_ratio']), 6)}，评级为 {three_way['grade']}。这个数字只能说明策略离“完美未来收益排序”非常远，不能说明策略本身没有价值。"
            ),
            Heading(1, "6. 工程实现：版本固化与标准化评估管线"),
            Paragraph(
                "为了避免“报告结论”和“默认配置”再次漂移，本周将版本注册、配置物化、标准化评估与分析入口全部整理为正式工程资产。当前已经可以基于统一注册表复现版本实验，并为单版本输出结构化的训练、事件、回测、分层分析、图表和报告。"
            ),
            TableBlock(
                headers=["Path", "Role"],
                rows=[
                    [str(ROOT / "configs" / "version_registry.yaml"), "集中维护 random 与 v1~v8 的正式版本定义"],
                    [str(ROOT / "configs" / "versions"), "存放物化后的 model/backtest 正式配置"],
                    [str(ROOT / "scripts" / "materialize_version_configs.py"), "根据注册表生成正式配置文件"],
                    [str(ROOT / "scripts" / "run_report_version_suite.py"), "批量运行当前数据版本的版本套件"],
                    [str(ROOT / "scripts" / "run_version_flip_analysis.py"), "生成 v1 相对 v7/v8 的分层翻转分析"],
                    [str(ROOT / "scripts" / "run_standardized_pipeline.py"), "单版本标准化评估入口，输出 metrics、analysis、charts 与 report"],
                ],
            ),
            TableBlock(
                headers=["Standardized Output", "Artifact"],
                rows=[
                    ["训练与预测", "preds/*.parquet, metrics.json"],
                    ["事件评估", "eval_stratified/summary.json"],
                    ["回测主结果", "backtest/summary.json"],
                    ["参考基准", "backtest/reference/*.json"],
                    ["分层分析", "standardized/analysis/*.csv"],
                    ["图表", "standardized/charts/*.png"],
                    ["报告", "standardized/report.md"],
                ],
            ),
            Heading(1, "7. 关键代码修改与验证"),
            TableBlock(
                headers=["File", "Change"],
                rows=[
                    [str(ROOT / "src" / "backtest" / "signal.py"), "抽出共享候选池构建逻辑，支持 strategy 与 reference 共用规则过滤"],
                    [str(ROOT / "src" / "backtest" / "benchmark.py"), "新增 non_prob 与 random_prob 参考线，并重构 oracle_event 口径"],
                    [str(ROOT / "src" / "backtest" / "report.py"), "新增 ranking_comparison 与额外 reference 输出写盘"],
                    [str(ROOT / "scripts" / "run_backtest.py"), "统一生成 strategy、non_prob、random_prob、y_div_10d 与 oracle_return 参考结果"],
                    [str(ROOT / "tests" / "test_backtest.py"), "补充共享候选池与 reference 构造测试，当前共 12 个测试用例"],
                ],
            ),
            TableBlock(
                headers=["Command", "Result"],
                rows=[
                    ["python -m pytest tests/test_backtest.py", "12 个测试全部通过"],
                    ["python scripts/run_backtest.py --run_id report_version_suite_current_data_20260323_v1 --backtest_cfg configs/versions/v1_backtest.yaml --split test", "回测成功，reference 结果已更新"],
                    ["python scripts/run_standardized_pipeline.py --version v1 --run_id report_version_suite_current_data_20260323_v1 --force_eval --force_backtest", "标准化产物存在且可复查"],
                ],
            ),
            Paragraph(
                f"当前 v1 的事件评估仍然保持较强识别能力：验证集 AUPRC 为 {pct(float(v1_metrics['training']['val_aucpr']))}，测试集事件命中率为 {pct(float(v1_metrics['event_eval']['event_level_summary']['hit_rate_events']))}。但 v7 在事件识别上更强，测试期事件命中率达到 {pct(float(v7_metrics['event_eval']['event_level_summary']['hit_rate_events']))}，回测却仍不如 v1；这进一步印证了“事件识别能力”和“交易收益排序能力”并不等价。"
            ),
            Paragraph(
                f"v8 的测试期事件命中率为 {pct(float(v8_metrics['event_eval']['event_level_summary']['hit_rate_events']))}，与 v1 接近，但测试期超额年化收益仍为 {pct(float(version_map['v8']['test_excess_annualized_return']))}。因此，当前数据版本下的主要问题不是‘是否能把事件猜出来’，而是‘在会发生事件的股票里，是否能按净收益做对排序’。"
            ),
            Heading(1, "8. 结果归档"),
            TableBlock(
                headers=["Artifact", "Description"],
                rows=[
                    [str(SUITE_DIR / "suite_results.csv"), "当前数据版本的版本套件总表"],
                    [str(SUITE_DIR / "current_data_experiment_report.md"), "当前数据版本实验报告"],
                    [str(FLIP_DIR / "summary.md"), "v1 相对 v7/v8 的翻转分析摘要"],
                    [str(V1_RUN_DIR / "backtest" / "reference" / "ranking_comparison.json"), "四路排序对比结果"],
                    [str(V1_RUN_DIR / "backtest" / "reference" / "three_way_comparison.json"), "random / strategy / oracle 三线比较结果"],
                    [str(RUNS_DIR / "report_version_suite_current_data_20260323_v1" / "standardized"), "v1 标准化评估产物"],
                    [str(RUNS_DIR / "report_version_suite_current_data_20260323_v7" / "standardized"), "v7 标准化评估产物"],
                    [str(RUNS_DIR / "report_version_suite_current_data_20260323_v8" / "standardized"), "v8 标准化评估产物"],
                ],
            ),
            Heading(1, "9. 本周总结与下周计划"),
            Heading(2, "9.1 本周总结"),
            BulletList(
                [
                    "在当前正确数据版本下完成了版本套件重跑，确认 v1 重新成为最优基线。",
                    "完成 v1 相对 v7/v8 的分层翻转分析，结论从“事件命中率更高”修正为“收益兑现更好”。",
                    "完成四路排序拆解，明确候选池价值与排序价值分别来自哪里。",
                    "完成 oracle 口径澄清，避免把 oracle_return 当作事件基准误读。",
                    "完成版本固化、标准化评估管线与测试补全，减少后续报告和代码口径漂移。",
                ]
            ),
            Heading(2, "9.2 下周计划"),
            BulletList(
                [
                    "把 run_report_version_suite.py 进一步收敛到标准化评估入口，减少重复执行链。",
                    "继续做 v1 在 2021 年相对落后的专项诊断，确认是否存在显著的时间漂移或成本结构变化。",
                    "在当前共享规则框架下补充更系统的成本敏感性分析，评估持有期、换手与 spread 对净收益的贡献。",
                    "将四路排序对比与 oracle 解释同步写回正式实验报告，作为之后版本复现的固定章节。",
                ]
            ),
        ]
    )

    return blocks


def main() -> None:
    blocks = build_blocks()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    render_docx(blocks)
    MARKDOWN_PATH.write_text(render_markdown(blocks), encoding="utf-8")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {MARKDOWN_PATH}")


if __name__ == "__main__":
    main()
